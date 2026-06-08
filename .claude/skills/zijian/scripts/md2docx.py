#!/usr/bin/env python3
"""
md2docx.py - 将自荐信Markdown转为Word文档

用法：python md2docx.py [input.md] [output.docx]
默认：读取 cover_letter_draft.md → 输出 cover_letter.docx
"""

import sys
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def parse_markdown(md_content):
    """解析Markdown内容，返回段落列表"""
    lines = md_content.split("\n")
    paragraphs = []

    for line in lines:
        stripped = line.strip()

        # 空行
        if not stripped:
            paragraphs.append("")
            continue

        # 检查链接 [text](url)
        link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'
        if re.search(link_pattern, stripped):
            # 保留带链接的文本（Word会自动识别超链接）
            paragraphs.append(stripped)
        else:
            paragraphs.append(stripped)

    return paragraphs


def create_word_document(paragraphs, output_file):
    """创建Word文档"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)  # 小四

    for text in paragraphs:
        if not text:
            # 空行
            doc.add_paragraph("")
            continue

        # 添加段落
        p = doc.add_paragraph()

        # 设置首行缩进
        p.paragraph_format.first_line_indent = Cm(0.74)  # 约2字符

        # 设置行距（1.5倍）
        p.paragraph_format.line_spacing = 1.5

        # 设置对齐方式（两端对齐）
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # 检查是否有链接
        link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'
        match = re.search(link_pattern, text)

        if match:
            # 有链接：分割文本和链接
            before = text[:match.start()]
            link_text = match.group(1)
            after = text[match.end():]

            # 添加链接前的文本
            if before:
                run = p.add_run(before)
                run.font.name = '宋体'
                run.font.size = Pt(12)

            # 添加超链接
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            # 创建超链接元素
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), '')

            # 简化处理：添加蓝色带下划线的文本
            run = p.add_run(link_text)
            run.font.name = '宋体'
            run.font.size = Pt(12)
            run.font.color.rgb = None  # 默认蓝色
            run.underline = True

            # 添加链接后的文本
            if after:
                run = p.add_run(after)
                run.font.name = '宋体'
                run.font.size = Pt(12)
        else:
            # 无链接：直接添加文本
            run = p.add_run(text)
            run.font.name = '宋体'
            run.font.size = Pt(12)

    # 保存文档
    doc.save(output_file)
    print(f"Done! Saved to {output_file}")


def main():
    # 解析命令行参数
    input_file = sys.argv[1] if len(sys.argv) > 1 else "cover_letter_draft.md"
    output_file = sys.argv[2] if len(sys.argv) > 2 else "cover_letter.docx"

    # 读取Markdown文件
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # 解析Markdown
    paragraphs = parse_markdown(content)

    # 生成Word文档
    create_word_document(paragraphs, output_file)


if __name__ == "__main__":
    main()
